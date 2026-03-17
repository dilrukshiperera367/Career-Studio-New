"""
Resume Parsing Service — Full deterministic pipeline.

No external AI APIs. Uses: regex, dictionaries, pdfplumber, python-docx, Tesseract, spaCy (local).

Pipeline: extract_text → clean_text → detect_sections → extract_entities → normalize → compute_derived
"""

import re
import math
import hashlib
import logging
from io import BytesIO
from collections import Counter
from datetime import date, datetime
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Heading dictionary for section detection (case-insensitive)
# ---------------------------------------------------------------------------

HEADING_MAP = {
    "experience": [
        "experience", "employment", "work history",
        "professional experience", "career history",
        "employment history",
    ],
    "education": [
        "education", "academic", "qualifications",
        "academic background", "academic qualifications",
        "degrees and certifications",
    ],
    "skills": [
        "skills", "technical skills", "core competencies",
        "tools", "technologies", "tools and technologies",
        "competencies",
    ],
    "projects": ["projects", "portfolio", "personal projects"],
    "certifications": [
        "certifications", "certificates", "licenses",
        "certifications and licenses", "licenses and certifications",
    ],
    "summary": [
        "summary", "profile", "objective",
        "professional summary", "professional profile", "about me",
    ],
}

DEGREE_KEYWORDS = [
    "bsc", "msc", "bachelor", "master", "phd", "university", "college",
    "diploma", "associate", "doctorate", "mba", "b.a.", "b.s.", "m.a.", "m.s.",
]

# ---------------------------------------------------------------------------
# Contact info patterns
# ---------------------------------------------------------------------------

CONTACT_PATTERNS = {
    "email": re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+"),
    "phone": re.compile(r"[\+]?[\d\s\-\(\)]{7,15}"),
    "linkedin": re.compile(r"linkedin\.com/in/[\w-]+", re.I),
    "github": re.compile(r"github\.com/[\w-]+", re.I),
    "portfolio": re.compile(r"https?://[\w.-]+\.[\w]{2,}[/\w.-]*"),
}

# ---------------------------------------------------------------------------
# Date range patterns
# ---------------------------------------------------------------------------

DATE_RANGE_PATTERNS = [
    # "Jan 2021 - Present", "January 2021 – Dec 2023"
    re.compile(
        r"(?P<start>(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4})"
        r"\s*[-–—]\s*"
        r"(?P<end>Present|Current|Now|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{4})",
        re.I,
    ),
    # "01/2021 - 12/2023"
    re.compile(
        r"(?P<start>\d{1,2}/\d{4})\s*[-–—]\s*(?P<end>Present|Current|\d{1,2}/\d{4})",
        re.I,
    ),
    # "2021 - 2023"
    re.compile(
        r"(?P<start>\d{4})\s*[-–—]\s*(?P<end>Present|Current|\d{4})",
        re.I,
    ),
]


# ===========================================================================
# 1. TEXT EXTRACTION
# ===========================================================================

def extract_text(file_content: bytes, file_type: str) -> str:
    """Extract raw text from resume file."""
    # ClamAV scan before processing (second-pass defence)
    try:
        from apps.shared.clamav import scan_bytes
        scan_result = scan_bytes(file_content)
        if not scan_result.is_clean:
            raise ValueError(f"File rejected by virus scanner: {scan_result.threat}")
    except ValueError:
        raise
    except Exception as _e:
        logger.warning("ClamAV scan skipped in parsing pipeline: %s", _e)

    MIN_TEXT_CHARS = 300

    if file_type in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"):
        text = _extract_docx(file_content)
    elif file_type in ("application/pdf", "pdf"):
        text = _extract_pdf(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

    # OCR fallback if too little text extracted
    if len(text.strip()) < MIN_TEXT_CHARS:
        try:
            text = _ocr_fallback(file_content, file_type)
        except Exception as e:
            logger.warning(f"OCR fallback failed: {e}")

    return text


def _extract_pdf(file_content: bytes) -> str:
    """Layout-aware PDF text extraction."""
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed")
        return ""

    full_text = []
    with pdfplumber.open(BytesIO(file_content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text(layout=True)
            if text:
                full_text.append(text)
    return "\n".join(full_text)


def _extract_docx(file_content: bytes) -> str:
    """Extract text from DOCX files."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed")
        return ""

    doc = Document(BytesIO(file_content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs)


def _ocr_fallback(file_content: bytes, file_type: str) -> str:
    """OCR fallback using Tesseract."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.error("pytesseract/Pillow not installed")
        return ""

    if file_type in ("application/pdf", "pdf"):
        try:
            import pdfplumber
            texts = []
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    img = page.to_image(resolution=300)
                    text = pytesseract.image_to_string(img.original)
                    texts.append(text)
            return "\n".join(texts)
        except Exception as e:
            logger.warning(f"PDF OCR failed: {e}")
            return ""
    else:
        img = Image.open(BytesIO(file_content))
        return pytesseract.image_to_string(img)


# ===========================================================================
# 2. TEXT CLEANING
# ===========================================================================

def clean_text(raw_text: str, page_count: Optional[int] = None) -> str:
    """Clean extracted text: normalize, remove artifacts, fix formatting."""
    text = raw_text

    # 1. Normalize newlines
    text = text.replace("\r\n", "\n")

    # 2. Remove repeated headers/footers
    lines = text.split("\n")
    line_freq = Counter(line.strip() for line in lines if line.strip())
    threshold = max(3, int((page_count or 10) * 0.3))
    repeated = {line for line, count in line_freq.items() if count >= threshold}
    lines = [l for l in lines if l.strip() not in repeated]
    text = "\n".join(lines)

    # 3. Fix hyphenation: "soft-\nware" → "software"
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # 4. Normalize bullets
    text = re.sub(r"^[\s]*[•·●○▪►▸⦿◦‣⁃]\s*", "* ", text, flags=re.MULTILINE)
    text = re.sub(r"^[\s]*[-–—]\s+", "* ", text, flags=re.MULTILINE)

    # 5. Collapse extra whitespace
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # 6. Strip page numbers
    text = re.sub(r"\n\s*Page\s+\d+\s*(of\s+\d+)?\s*\n", "\n", text, flags=re.I)

    return text.strip()


# ===========================================================================
# 3. SECTION DETECTION
# ===========================================================================

def detect_sections(clean_text_str: str) -> dict:
    """Split resume into labeled sections using heading detection + fallback."""
    lines = clean_text_str.split("\n")
    heading_indices = []

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue

        if len(stripped) < 40:
            normalized = stripped.lower().rstrip(":").strip()
            for section, keywords in HEADING_MAP.items():
                if normalized in keywords:
                    heading_indices.append((i, section))
                    break
            else:
                # Detect mostly-uppercase headings
                if len(stripped) > 3:
                    alpha_chars = [c for c in stripped if c.isalpha()]
                    if alpha_chars and sum(1 for c in alpha_chars if c.isupper()) / len(alpha_chars) > 0.7:
                        normalized_upper = stripped.lower().rstrip(":").strip()
                        for section, keywords in HEADING_MAP.items():
                            if normalized_upper in keywords:
                                heading_indices.append((i, section))
                                break

    # Build section ranges
    sections = {}
    for idx, (line_i, section) in enumerate(heading_indices):
        end_i = heading_indices[idx + 1][0] if idx + 1 < len(heading_indices) else len(lines)
        sections[section] = "\n".join(lines[line_i + 1: end_i])

    # Collect text before first heading as "header"
    if heading_indices:
        sections["header"] = "\n".join(lines[: heading_indices[0][0]])

    # FALLBACK: if no headings found, detect by patterns
    if not sections:
        sections = _detect_sections_by_patterns(lines)

    return sections


def _detect_sections_by_patterns(lines: list) -> dict:
    """Fallback section detection when no explicit headings found."""
    date_pattern = re.compile(r"\b\d{4}\s*[-–—]\s*(\d{4}|present|current)\b", re.I)

    education_lines = []
    experience_lines = []
    other_lines = []

    for line in lines:
        lower = line.lower()
        if any(kw in lower for kw in DEGREE_KEYWORDS):
            education_lines.append(line)
        elif date_pattern.search(line):
            experience_lines.append(line)
        else:
            other_lines.append(line)

    return {
        "education": "\n".join(education_lines),
        "experience": "\n".join(experience_lines),
        "other": "\n".join(other_lines),
    }


# ===========================================================================
# 4. ENTITY EXTRACTION
# ===========================================================================

def extract_contact_info(text: str) -> dict:
    """Extract emails, phones, LinkedIn, GitHub from header/full text."""
    result = {"emails": [], "phones": [], "urls": {}}

    for email in CONTACT_PATTERNS["email"].findall(text):
        if email not in result["emails"]:
            result["emails"].append(email)

    for phone in CONTACT_PATTERNS["phone"].findall(text):
        cleaned = re.sub(r"[^\d+]", "", phone)
        if len(cleaned) >= 7 and cleaned not in [p for p in result["phones"]]:
            result["phones"].append(phone.strip())

    for linkedin in CONTACT_PATTERNS["linkedin"].findall(text):
        result["urls"]["linkedin"] = linkedin

    for github in CONTACT_PATTERNS["github"].findall(text):
        result["urls"]["github"] = github

    return result


def extract_date_ranges(text: str) -> list:
    """Find all date ranges in text."""
    ranges = []
    for pattern in DATE_RANGE_PATTERNS:
        for match in pattern.finditer(text):
            ranges.append({
                "start": match.group("start"),
                "end": match.group("end"),
                "span": match.span(),
            })
    return ranges


def extract_experience_blocks(experience_text: str) -> list:
    """Heuristic: find date-range anchors, look ±3 lines for title/company."""
    lines = experience_text.split("\n")
    blocks = []
    used_lines = set()

    for i, line in enumerate(lines):
        if i in used_lines:
            continue

        date_ranges = extract_date_ranges(line)
        if not date_ranges:
            continue

        dr = date_ranges[0]
        window_start = max(0, i - 3)
        window_end = min(len(lines), i + 4)
        window = lines[window_start:window_end]

        title, company = _extract_title_company(window)

        blocks.append({
            "company": company,
            "title": title,
            "start": dr["start"],
            "end": dr["end"],
            "raw_block": "\n".join(window),
        })
        used_lines.update(range(window_start, window_end))

    return blocks


def _extract_title_company(window_lines: list) -> tuple:
    """Try to extract title and company from an experience window."""
    # Pattern: "Title at Company"
    at_pattern = re.compile(
        r"(?P<title>[A-Z][\w\s&/,]+?)\s*(?:at|@|[-–—|])\s*(?P<company>[A-Z][\w\s&.,]+)",
        re.I,
    )
    for line in window_lines:
        match = at_pattern.search(line)
        if match:
            return match.group("title").strip(), match.group("company").strip()

    # Fallback: first non-date, non-empty line is title, next is company
    non_empty = [l.strip() for l in window_lines if l.strip() and not extract_date_ranges(l)]
    title = non_empty[0] if len(non_empty) > 0 else ""
    company = non_empty[1] if len(non_empty) > 1 else ""
    return title, company


def extract_skills(clean_text_str: str, taxonomy: dict) -> list:
    """
    Match n-grams (1-3 words) against skill alias map.

    taxonomy: {alias_normalized: {"skill_id": "...", "canonical_name": "..."}}
    """
    text_lower = clean_text_str.lower()
    words = re.findall(r"\b\w+\b", text_lower)

    found = {}

    for n in (3, 2, 1):
        for i in range(len(words) - n + 1):
            ngram = " ".join(words[i: i + n])
            ngram_norm = re.sub(r"[^\w\s]", "", ngram).strip()
            ngram_norm = re.sub(r"\s+", " ", ngram_norm)

            if ngram_norm in taxonomy:
                skill = taxonomy[ngram_norm]
                sid = skill["skill_id"]
                if sid not in found:
                    found[sid] = {
                        "skill_id": sid,
                        "canonical_name": skill["canonical_name"],
                        "confidence": 1.0,
                        "evidence": [],
                    }
                if ngram not in found[sid]["evidence"]:
                    found[sid]["evidence"].append(ngram)

    return list(found.values())


# ===========================================================================
# 5. NORMALIZATION + DERIVED FIELDS
# ===========================================================================

def normalize_title(raw_title: str, alias_map: dict) -> str:
    """Map raw title to canonical form if alias exists."""
    normalized = raw_title.lower().strip()
    return alias_map.get(normalized, raw_title)


def normalize_location(raw_location: str, location_map: dict) -> dict:
    """Map raw location to structured city/country."""
    normalized = raw_location.lower().strip()
    match = location_map.get(normalized)
    return match if match else {"raw": raw_location, "city": None, "country_code": None}


def parse_fuzzy_date(date_str: str) -> Optional[date]:
    """Parse a fuzzy date string into a date object."""
    if not date_str:
        return None
    date_str = date_str.strip()
    if date_str.lower() in ("present", "current", "now"):
        return date.today()

    try:
        from dateutil.parser import parse as dateutil_parse
        return dateutil_parse(date_str, fuzzy=True).date()
    except Exception:
        pass

    # Try year-only
    year_match = re.match(r"(\d{4})", date_str)
    if year_match:
        return date(int(year_match.group(1)), 1, 1)

    return None


def compute_derived_fields(experiences: list, skills: list) -> dict:
    """Compute total_years, recent_title/company, recency_score."""
    # Total experience years (merge overlapping intervals)
    intervals = []
    for exp in experiences:
        start = parse_fuzzy_date(exp.get("start", ""))
        end_str = exp.get("end", "")
        end = parse_fuzzy_date(end_str) if end_str and end_str.lower() not in ("present", "current", "now") else date.today()
        if start and end and end >= start:
            intervals.append((start, end))

    total_years = _sum_non_overlapping_intervals(intervals)

    # Most recent title/company
    sorted_exp = sorted(
        experiences,
        key=lambda e: parse_fuzzy_date(e.get("end", "")) or date.today(),
        reverse=True,
    )
    recent_title = sorted_exp[0].get("title", "") if sorted_exp else ""
    recent_company = sorted_exp[0].get("company", "") if sorted_exp else ""

    # Recency score
    recent_role_text = (sorted_exp[0].get("raw_block", "") if sorted_exp else "").lower()
    skill_recency = {}

    for skill in skills:
        evidence_strs = [ev.lower() for ev in skill.get("evidence", [])]
        if any(ev in recent_role_text for ev in evidence_strs):
            skill_recency[skill["skill_id"]] = 1.0
        else:
            years_since = _estimate_years_since_skill(skill, experiences)
            decay = 0.35  # from spec
            skill_recency[skill["skill_id"]] = max(0.0, min(1.0, math.exp(-decay * years_since)))

    recency_score = (
        sum(skill_recency.values()) / max(len(skill_recency), 1)
        if skill_recency else 0.5
    )

    return {
        "total_experience_years": round(total_years, 1),
        "most_recent_title": recent_title,
        "most_recent_company": recent_company,
        "recency_score": round(recency_score, 4),
    }


def _sum_non_overlapping_intervals(intervals: list) -> float:
    """Merge overlapping date intervals and return total years."""
    if not intervals:
        return 0.0
    sorted_iv = sorted(intervals)
    merged = [sorted_iv[0]]
    for start, end in sorted_iv[1:]:
        if start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return sum((end - start).days / 365.25 for start, end in merged)


def _estimate_years_since_skill(skill: dict, experiences: list) -> float:
    """Estimate how many years since a skill was last used in a role."""
    evidence_strs = [ev.lower() for ev in skill.get("evidence", [])]
    best_end = None

    for exp in experiences:
        raw = (exp.get("raw_block", "") + " " + exp.get("title", "")).lower()
        if any(ev in raw for ev in evidence_strs):
            end_date = parse_fuzzy_date(exp.get("end", ""))
            if end_date and (best_end is None or end_date > best_end):
                best_end = end_date

    if best_end is None:
        return 5.0  # default penalty if we can't determine

    return max(0.0, (date.today() - best_end).days / 365.25)


# ===========================================================================
# 6. FULL PIPELINE ORCHESTRATOR
# ===========================================================================

def parse_resume_full(file_content: bytes, file_type: str, taxonomy: dict) -> dict:
    """Run the complete parsing pipeline. Returns parsed_json."""
    # Step 1: Extract
    raw_text = extract_text(file_content, file_type)

    # Step 2: Clean
    cleaned = clean_text(raw_text)

    # Step 3: Detect sections
    sections = detect_sections(cleaned)

    # Step 4: Extract entities
    contact_info = extract_contact_info(sections.get("header", cleaned[:500]))

    experience_text = sections.get("experience", "")
    experience_blocks = extract_experience_blocks(experience_text)

    skills = extract_skills(cleaned, taxonomy)

    # Step 5: Compute derived fields
    derived = compute_derived_fields(experience_blocks, skills)

    # Build parsed JSON
    parsed = {
        "identities": contact_info,
        "experience": experience_blocks,
        "education": [],  # simplified — section text stored
        "skills": skills,
        "sections": {k: v[:500] for k, v in sections.items()},  # truncated for storage
        "derived": derived,
    }

    return {
        "raw_text": raw_text,
        "clean_text": cleaned,
        "parsed_json": parsed,
    }
